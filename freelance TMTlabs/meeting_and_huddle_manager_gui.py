import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import tkinter.font as tkfont
import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
import openpyxl
from openpyxl.styles import Font, PatternFill
import pickle
import matplotlib
matplotlib.use('TkAgg')
from matplotlib.figure import Figure
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

class NurseManagerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Meeting Agenda and Huddle Tool")
        # desired default window size (will be adjusted to fit the screen)
        default_w, default_h = 1500, 900
        # center horizontally and position just below the top of the screen
        try:
            self.root.update_idletasks()
            screen_w = self.root.winfo_screenwidth()
            screen_h = self.root.winfo_screenheight()
            w = min(default_w, max(200, screen_w - 40))
            h = min(default_h, max(200, screen_h - 20))
            x = max(0, (screen_w - w) // 2)
            # position flush with the top of the screen (outer window at y=0)
            y_offset = 0
            # apply geometry and force an update so the window manager honors placement
            try:
                self.root.geometry(f"{w}x{h}+{x}+{y_offset}")
                self.root.update()
                # ensure the window is lifted briefly so placement is respected
                self.root.lift()
                self.root.attributes('-topmost', True)
                # remove topmost soon after to restore normal stacking
                self.root.after(200, lambda: self.root.attributes('-topmost', False))
            except Exception:
                # fallback to prior behavior if attributes aren't supported
                try:
                    self.root.geometry(f"{w}x{h}+{x}+{y_offset}")
                except Exception:
                    pass
        except Exception:
            # fallback to a fixed geometry if anything goes wrong
            try:
                self.root.geometry("1500x900")
            except Exception:
                pass
        # set a larger default font
        self.font_name = "Segoe UI"
        self.font_size = 14
        try:
            # update the default Tk font so widgets inherit it
            default_font = tkfont.nametofont("TkDefaultFont")
            default_font.configure(family=self.font_name, size=self.font_size)
            self.default_font = default_font
            style = ttk.Style()
            style.configure('TNotebook.Tab', font=(self.font_name, self.font_size))
            # make scrollbar more visible with higher contrast
            style.configure('Big.Vertical.TScrollbar', troughcolor='#ffffff', background='#4a4a4a', troughrelief='flat')
            style.map('Big.Vertical.TScrollbar', background=[('active', '#2f2f2f'), ('!active', '#4a4a4a')])
            style.configure('TNotebook', padding=6)
        except Exception:
            self.default_font = None
        
        # Metrics list for dropdown
        self.metrics = [
            "behavioral health ptsd screen",
            "behavioral health depression screen",
            "behavioral health alcohol screen(audit-c)",
            "hypertension bp less than 140/90",
            "diabetes- A1c >9 or not done (poor control) in past year",
            "diabetes- bp less than 140/90",
            "diabetes -retinal exam",
            "diabetes - kidney health",
            "statin therapy diabetes",
            "statin therapy ischemic heart disease",
            "prevention- tobacco counselling- timely counselling",
            "prevention cervical cancer screening",
            "prevention breast cancer screening (40-74)",
            "prevention breast cancer screen (50-74)"
        ]
        
        # numeric keys used for validation/highlighting
        self.numeric_keys = [
            'total_discharges', 'calls_completed', 'completion_rate', 'missed_calls',
            'dm_total_patients', 'dm_a1c_updated', 'dm_microalbuminuria',
            'total_pact_teams', 'red_teams_count', 'charts_audited', 'outliers_count'
        ]

        self.data = {}
        # validation marker labels for top-level fields
        self.validation_markers = {}
        self.current_file = None
        self.discipline_issues = []
        self.pdsa_projects = []
        self.red_teams = []
        self.create_menu()
        self.create_widgets()
        
    def create_menu(self):
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="New Form", command=self.new_form)
        file_menu.add_command(label="Open", command=self.open_form)
        file_menu.add_command(label="Save", command=self.save_form)
        file_menu.add_command(label="Save As", command=self.save_form_as)
        file_menu.add_separator()
        file_menu.add_command(label="Export to Word", command=self.export_to_word)
        file_menu.add_command(label="Export to Excel", command=self.export_to_excel)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.root.quit)
        
        help_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Help", menu=help_menu)
        help_menu.add_command(label="About", command=self.show_about)
    
    def create_widgets(self):
        # Main notebook for tabs
        notebook = ttk.Notebook(self.root)
        notebook.pack(fill="both", expand=True, padx=5, pady=5)
        
        # Create tabs
        self.weekly_tab = self.create_weekly_tab(notebook)
        self.monthly_tab = self.create_monthly_tab(notebook)
        self.pdsa_tab = self.create_pdsa_tab(notebook)
        self.support_tab = self.create_support_tab(notebook)
        
    def create_weekly_tab(self, notebook):
        frame = ttk.Frame(notebook)
        notebook.add(frame, text="Weekly")
        
        # Create scrollbar
        canvas = tk.Canvas(frame)
        scrollbar = tk.Scrollbar(frame, orient="vertical", command=canvas.yview, width=30, bg='#4a4a4a', troughcolor='#ffffff', activebackground='#2f2f2f')
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Header info
        header = self.create_form_section(scrollable_frame, "Huddle Information")
        self.service_unit = self.add_text_field(header, "Service/Unit:", "service_unit")
        self.nurse_manager = self.add_text_field(header, "Nurse Manager:", "nurse_manager")
        self.week_of = self.add_text_field(header, "Week of (Mon-Sun):", "week_of")
        self.huddle_date = self.add_text_field(header, "Huddle Date:", "huddle_date")
        self.participants = self.add_text_field(header, "Participants:", "participants", height=3)
        
        # 48 Hour Discharge Calls
        discharge_frame = self.create_form_section(scrollable_frame, "48 Hour Discharge Calls")
        self.total_discharges = self.add_text_field(discharge_frame, "Total Discharges Last Week:", "total_discharges")
        self.calls_completed = self.add_text_field(discharge_frame, "Total 48 Hour Calls Completed:", "calls_completed")
        self.completion_rate = self.add_text_field(discharge_frame, "Completion Rate (%):", "completion_rate")
        self.missed_calls = self.add_text_field(discharge_frame, "Missed Calls (number):", "missed_calls")
        ttk.Button(discharge_frame, text="Show Discharge Chart", command=self.show_discharge_chart).grid(row=len(discharge_frame.grid_slaves()), column=1, sticky='e', padx=8, pady=6)

        # Missed calls reasons (place under the chart button)
        mrow = len([w for w in discharge_frame.winfo_children()])
        ttk.Label(discharge_frame, text="Missed Call Reasons:").grid(row=mrow, column=0, sticky="nw", pady=(10,5))
        self.no_phone = self.add_text_field(discharge_frame, "No valid phone number:", "no_phone")
        self.unable_reach = self.add_text_field(discharge_frame, "Unable to reach after 2 attempts:", "unable_reach")
        self.staffing_shortage = self.add_text_field(discharge_frame, "Staffing shortage:", "staffing_shortage")
        self.workflow_issue = self.add_text_field(discharge_frame, "Workflow/process issue:", "workflow_issue")
        self.other_reason = self.add_text_field(discharge_frame, "Other (specify):", "other_reason")
        
        # Actions taken
        actions_frame = self.create_form_section(scrollable_frame, "Actions Taken This Week")
        self.action_staffing = self.add_checkbox(actions_frame, "Adjusted staffing/assignment", "action_staffing")
        self.action_workflow = self.add_checkbox(actions_frame, "Clarified workflow/standard work", "action_workflow")
        self.action_patient_list = self.add_checkbox(actions_frame, "Updated patient list/source", "action_patient_list")
        self.action_education = self.add_checkbox(actions_frame, "Provided staff education", "action_education")
        self.action_it = self.add_checkbox(actions_frame, "Escalated IT issue", "action_it")
        self.action_other = self.add_text_field(actions_frame, "Other:", "action_other")
        
        # Issues to escalate
        issues_frame = self.create_form_section(scrollable_frame, "Issues to Prioritize/Support Needed")
        self.issues_escalate = self.add_text_field(issues_frame, "", "issues_escalate", height=4)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        return frame

    def add_discipline_issue(self, prefill=None):
        """Add a discipline issue entry to the UI"""
        idx = len(self.discipline_issues) + 1
        frame = ttk.LabelFrame(self.discipline_container, text=f"Issue #{idx}", padding=6)
        frame.pack(fill="x", padx=5, pady=5)

        fields = {}
        def mk(label, key, height=1):
            row = len([w for w in frame.grid_slaves()])
            ttk.Label(frame, text=label).grid(row=row, column=0, sticky="ne", padx=5, pady=3)
            if height == 1:
                w = ttk.Entry(frame, width=50)
            else:
                w = tk.Text(frame, height=height, width=50)
            w.grid(row=row, column=1, sticky="ew", padx=5, pady=3)
            fields[key] = w
            if prefill and key in prefill:
                if isinstance(w, tk.Text):
                    w.insert("1.0", prefill[key])
                else:
                    w.insert(0, prefill[key])

        mk("Issue Description (no names):", "description", height=2)
        mk("Coaching - Staff Initials:", "coaching_initials")
        mk("Coaching - Date Started (MM/DD/YYYY):", "coaching_date")
        mk("Coaching - Status:", "coaching_status")
        mk("Written Counseling - Staff Initials:", "written_initials")
        mk("Written Counseling - Date (MM/DD/YYYY):", "written_date")
        mk("Written Counseling - Follow-up Due (MM/DD/YYYY):", "written_followup")
        mk("Performance Improvement Plan - Staff Initials:", "pip_initials")
        mk("Performance Improvement Plan - Start Date (MM/DD/YYYY):", "pip_start")
        mk("Performance Improvement Plan - Review Date (MM/DD/YYYY):", "pip_review")
        mk("HR/Union/Legal Deadlines:", "deadlines", height=2)
        mk("Support Needed (HR documentation, meeting, etc):", "support", height=2)

        remove_btn = ttk.Button(frame, text="Remove", command=lambda f=frame: self.remove_discipline_issue(f))
        rrow = len([w for w in frame.grid_slaves()])
        remove_btn.grid(row=rrow, column=1, sticky="e", padx=5, pady=5)

        self.discipline_issues.append((frame, fields))

    def remove_discipline_issue(self, frame):
        """Remove a discipline issue UI entry"""
        for i, (frm, fields) in enumerate(self.discipline_issues):
            if frm == frame:
                frm.destroy()
                del self.discipline_issues[i]
                break
        for idx, (frm, _) in enumerate(self.discipline_issues, start=1):
            frm.config(text=f"Issue #{idx}")

    def add_pdsa_project(self, prefill=None):
        """Add a PDSA project entry to the UI"""
        idx = len(self.pdsa_projects) + 1
        frame = ttk.LabelFrame(self.pdsa_container, text=f"PDSA Project #{idx}", padding=6)
        frame.pack(fill="x", padx=5, pady=5)

        fields = {}
        def mk(label, key, height=1, width=50):
            row = len([w for w in frame.grid_slaves()])
            ttk.Label(frame, text=label).grid(row=row, column=0, sticky="ne", padx=5, pady=3)
            if height == 1:
                w = ttk.Entry(frame, width=width)
            else:
                w = tk.Text(frame, height=height, width=width)
            w.grid(row=row, column=1, sticky="ew", padx=5, pady=3)
            fields[key] = w
            if prefill and key in prefill:
                if isinstance(w, tk.Text):
                    w.insert("1.0", prefill[key])
                else:
                    w.insert(0, prefill[key])
            # if this is a metric numeric field, add a grid marker and bind validation
            if key.endswith('_value') or key.endswith('_goal'):
                marker = tk.Label(frame, text='', fg='red')
                try:
                    marker.configure(font=(self.font_name, max(self.font_size-2, 10)))
                except Exception:
                    pass
                marker.grid(row=row, column=2, sticky='w', padx=4)

                def on_focus_out_pdsa(event, widget=w, m=marker):
                    val = widget.get() if not isinstance(widget, tk.Text) else widget.get('1.0','end-1c')
                    if val and self._parse_number(val) is None:
                        m.config(text='✱ invalid number')
                    else:
                        m.config(text='')

                try:
                    w.bind('<FocusOut>', on_focus_out_pdsa)
                except Exception:
                    pass

        mk("Project Name:", "name")
        mk("Focus Area:", "focus")
        phase_var = tk.StringVar()
        fields['phase'] = phase_var
        phase_frame = ttk.Frame(frame)
        prow = len([w for w in frame.grid_slaves()])
        phase_frame.grid(row=prow, column=0, columnspan=2, sticky="ew", pady=3)
        ttk.Label(phase_frame, text="Current Phase:").pack(side="left")
        for phase in ["Plan", "Do", "Study", "Act"]:
            ttk.Radiobutton(phase_frame, text=phase, variable=phase_var, value=phase).pack(side="left", padx=5)
        if prefill and 'phase' in prefill:
            phase_var.set(prefill.get('phase',''))

        for j in range(1, 4):
            mk(f"Metric {j}:", f"metric{j}")
            mk(f"  Value:", f"metric{j}_value", width=20)
            mk(f"  Goal:", f"metric{j}_goal", width=20)

        mk("Key Learning from This Month's Data:", "learning", height=2)
        mk("Next Test/Change Before Next Huddle:", "next_test", height=2)

        remove_btn = ttk.Button(frame, text="Remove", command=lambda f=frame: self.remove_pdsa_project(f))
        rrow = len([w for w in frame.grid_slaves()])
        remove_btn.grid(row=rrow, column=1, sticky="e", padx=5, pady=5)

        self.pdsa_projects.append((frame, fields))

    def remove_pdsa_project(self, frame):
        """Remove a PDSA project UI entry"""
        for i, (frm, fields) in enumerate(self.pdsa_projects):
            if frm == frame:
                frm.destroy()
                del self.pdsa_projects[i]
                break
        for idx, (frm, _) in enumerate(self.pdsa_projects, start=1):
            frm.config(text=f"PDSA Project #{idx}")

    def add_red_team(self, prefill=None):
        """Add a red team entry with metrics and barriers"""
        idx = len(self.red_teams) + 1
        frame = ttk.LabelFrame(self.red_container, text=f"Team with red metrics #{idx}", padding=6)
        frame.pack(fill="x", padx=5, pady=5)

        info = {'metrics': [], 'barriers': [], 'name': None}

        # Team name
        ttk.Label(frame, text="Team Name/Number:").grid(row=0, column=0, sticky="ne", padx=5, pady=3)
        name_entry = ttk.Entry(frame, width=40)
        name_entry.grid(row=0, column=1, sticky="w", padx=5, pady=3)
        info['name'] = name_entry
        if prefill and 'name' in prefill:
            name_entry.insert(0, prefill.get('name',''))

        # Metrics container (start with one metric row)
        metrics_frame = ttk.Frame(frame)
        metrics_frame.grid(row=1, column=0, columnspan=2, sticky='ew', padx=5, pady=3)
        ttk.Label(metrics_frame, text="Metrics (select metric, current status, target goal)").pack(side='left')

        # store metrics_frame in info so helper can add rows later
        info['metrics_frame'] = metrics_frame

        # create one default metric row
        self._create_metric_row_in_frame(metrics_frame, info)

        def remove_metric_row(row_frame, team_info):
            for i, m in enumerate(team_info['metrics']):
                if m['frame'] == row_frame:
                    try:
                        m['frame'].destroy()
                    except Exception:
                        pass
                    del team_info['metrics'][i]
                    break

        # Barriers with per-barrier support field
        barriers_frame = ttk.LabelFrame(frame, text="Barriers Blocking Improvement Opportunities (check all that apply)")
        barriers_frame.grid(row=2, column=0, columnspan=2, sticky='ew', padx=5, pady=3)
        barrier_labels = [
            'data/reporting issue',
            'staffing/capacity',
            'provider engagement',
            'patient factors (no shows/refusals/social)',
            'workflow unclear/too complex',
            'informatics/template limitation',
            'other'
        ]
        for i, bl in enumerate(barrier_labels):
            var = tk.BooleanVar()
            chk = ttk.Checkbutton(barriers_frame, text=bl, variable=var)
            chk.grid(row=i, column=0, sticky='w', padx=5, pady=2)
            support = ttk.Entry(barriers_frame, width=60)
            support.grid(row=i, column=1, sticky='w', padx=5, pady=2)
            # disable support unless checked
            def make_trace(v, s):
                def trace(*args):
                    if v.get():
                        s.config(state='normal')
                    else:
                        s.delete(0, 'end')
                        s.config(state='disabled')
                return trace
            var.trace_add('write', make_trace(var, support))
            # set initial disabled
            support.config(state='disabled')
            info['barriers'].append({'label': bl, 'var': var, 'support': support})

        # prefill metrics and barriers if provided
        if prefill:
            for m in prefill.get('metrics', []):
                self._create_metric_row_in_frame(metrics_frame, info, m_prefill=m)
            # set barrier checks and support
            for b in prefill.get('barriers', []):
                for entry in info['barriers']:
                    if entry['label'] == b.get('name'):
                        if b.get('checked'):
                            entry['var'].set(True)
                            entry['support'].config(state='normal')
                            entry['support'].insert(0, b.get('support',''))

        # chart and remove team buttons
        btn_frame = ttk.Frame(frame)
        btn_frame.grid(row=99, column=0, columnspan=2, sticky='e', padx=5, pady=5)
        ttk.Button(btn_frame, text="Show Team Chart", command=lambda info=info: self.show_team_chart(info)).pack(side='right', padx=4)
        ttk.Button(btn_frame, text="Add Multiple Metrics", command=lambda info=info: self.open_add_metrics_dialog(info)).pack(side='right', padx=4)
        remove_btn = ttk.Button(btn_frame, text="Remove Team", command=lambda f=frame: self.remove_red_team(f))
        remove_btn.pack(side='right', padx=4)

        self.red_teams.append((frame, info))

    def remove_red_team(self, frame):
        for i, (frm, info) in enumerate(self.red_teams):
            if frm == frame:
                try:
                    frm.destroy()
                except Exception:
                    pass
                del self.red_teams[i]
                break
        for idx, (frm, _) in enumerate(self.red_teams, start=1):
            frm.config(text=f"Team With Red Metrics #{idx}")

    def _create_metric_row_in_frame(self, metrics_frame, team_info, m_prefill=None):
        """Create a single metric row inside the given metrics_frame and register it on team_info."""
        row = ttk.Frame(metrics_frame)
        row.pack(fill='x', padx=5, pady=2)
        combo = ttk.Combobox(row, values=self.metrics, width=40, state='readonly')
        combo.pack(side='left', padx=5)

        # Current field with label above
        cur_frame = ttk.Frame(row)
        cur_frame.pack(side='left', padx=5)
        ttk.Label(cur_frame, text='Current').pack(side='top')
        current = ttk.Entry(cur_frame, width=20)
        current.pack(side='top')
        cur_marker = tk.Label(cur_frame, text='', fg='red')
        try:
            cur_marker.configure(font=(self.font_name, max(self.font_size-4, 9)))
        except Exception:
            pass
        cur_marker.pack(side='left', padx=(6,0))

        # Goal field with label above
        goal_frame = ttk.Frame(row)
        goal_frame.pack(side='left', padx=5)
        ttk.Label(goal_frame, text='Goal').pack(side='top')
        goal = ttk.Entry(goal_frame, width=20)
        goal.pack(side='top')
        goal_marker = tk.Label(goal_frame, text='', fg='red')
        try:
            goal_marker.configure(font=(self.font_name, max(self.font_size-4, 9)))
        except Exception:
            pass
        goal_marker.pack(side='left', padx=(6,0))

        def remove_metric_row_local():
            for i, m in enumerate(team_info['metrics']):
                if m['frame'] == row:
                    try:
                        row.destroy()
                    except Exception:
                        pass
                    del team_info['metrics'][i]
                    break

        rem = ttk.Button(row, text='Remove', command=remove_metric_row_local)
        rem.pack(side='left', padx=5)

        m = {'frame': row, 'combo': combo, 'current': current, 'goal': goal}
        team_info.setdefault('metrics', []).append(m)
        if m_prefill:
            combo.set(m_prefill.get('metric',''))
            current.insert(0, m_prefill.get('current',''))
            goal.insert(0, m_prefill.get('goal',''))

        # bind focus-out validation for metric current/goal
        def make_validate(widget, marker):
            def on_focus_out(e):
                val = widget.get()
                if val and self._parse_number(val) is None:
                    marker.config(text='✱')
                else:
                    marker.config(text='')
            return on_focus_out

        try:
            current.bind('<FocusOut>', make_validate(current, cur_marker))
            goal.bind('<FocusOut>', make_validate(goal, goal_marker))
        except Exception:
            pass

        return m

    def open_add_metrics_dialog(self, info):
        """Open a dialog to multi-select metrics from the predefined list and add them to the team."""
        # ensure metrics_frame is available
        metrics_frame = info.get('metrics_frame')
        if metrics_frame is None:
            messagebox.showwarning('No Metrics Area', 'Cannot add metrics to this team (metrics area unavailable).')
            return

        existing = set()
        for m in info.get('metrics', []):
            try:
                val = m.get('combo').get() if m.get('combo') else ''
                if val:
                    existing.add(val)
            except Exception:
                pass

        win = tk.Toplevel(self.root)
        win.title('Add Multiple Metrics')
        win.geometry('520x620')

        ttk.Label(win, text='Select metrics to add for this team (multi-select)').pack(anchor='w', padx=10, pady=(10,4))
        lb = tk.Listbox(win, selectmode='multiple', width=70, height=18)
        lb.pack(padx=10, pady=4, fill='both', expand=True)
        # populate listbox
        for i, metric in enumerate(self.metrics):
            lb.insert('end', metric)
            # preselect metrics not already present? leave unselected; user chooses

        btn_frame = ttk.Frame(win)
        btn_frame.pack(fill='x', padx=10, pady=8)

        def do_add():
            sel = lb.curselection()
            added = 0
            for idx in sel:
                metric = lb.get(idx)
                if metric in existing:
                    continue
                self._create_metric_row_in_frame(metrics_frame, info, m_prefill={'metric': metric})
                added += 1
            if added == 0:
                messagebox.showinfo('No New Metrics', 'No new metrics were added (selected metrics may already exist).')
            win.destroy()

        ttk.Button(btn_frame, text='Add Selected', command=do_add).pack(side='right', padx=6)
        ttk.Button(btn_frame, text='Cancel', command=win.destroy).pack(side='right')

    def _parse_number(self, s):
        try:
            if s is None:
                return None
            if isinstance(s, (int, float)):
                return float(s)
            st = str(s).strip()
            if st.endswith('%'):
                st = st[:-1]
            st = st.replace(',', '')
            return float(st)
        except Exception:
            return None

    def _format_number_for_display(self, s):
        """Return a cleaned, human-readable numeric string or original if not parseable."""
        n = self._parse_number(s)
        if n is None:
            return str(s)
        # if integer-like, show as int
        if abs(n - int(n)) < 1e-9:
            return str(int(n))
        return str(n)

    def _collect_numeric_issues(self):
        """Check known numeric fields and red-team metric fields for parseable numbers.
        Returns list of (label, widget, raw_value) for problematic fields.
        """
        issues = []
        # fields expected numeric in self.data
        numeric_keys = [
            'total_discharges', 'calls_completed', 'completion_rate', 'missed_calls',
            'dm_total_patients', 'dm_a1c_updated', 'dm_microalbuminuria',
            'total_pact_teams', 'red_teams_count', 'charts_audited', 'outliers_count'
        ]
        for key in numeric_keys:
            if key in self.data:
                w = self.data[key]
                val = self.get_field_value(w)
                if val is None or str(val).strip() == '':
                    continue
                if self._parse_number(val) is None:
                    issues.append((key, w, val))

        # PDSA project metric values
        for _, fields in self.pdsa_projects:
            for j in range(1, 4):
                for suffix in ('_value', '_goal'):
                    k = f'metric{j}{suffix}'
                    if k in fields:
                        w = fields[k]
                        val = w.get() if not isinstance(w, tk.Text) else w.get('1.0','end-1c')
                        if val and self._parse_number(val) is None:
                            issues.append((k, w, val))

        # Red team metric values
        for _, info in self.red_teams:
            for m in info.get('metrics', []):
                cur = m.get('current')
                goal = m.get('goal')
                if cur:
                    v = cur.get()
                    if v and self._parse_number(v) is None:
                        issues.append((f"{m.get('combo').get() or 'metric_current'}", cur, v))
                if goal:
                    v = goal.get()
                    if v and self._parse_number(v) is None:
                        issues.append((f"{m.get('combo').get() or 'metric_goal'}", goal, v))

        return issues

    def _attempt_autoclean_numeric_issues(self, issues):
        """Try to autoclean numeric issues using _parse_number and write back cleaned value."""
        cleaned = []
        for key, widget, raw in issues:
            parsed = self._parse_number(raw)
            if parsed is not None:
                # write back a cleaned string
                clean_str = self._format_number_for_display(parsed)
                try:
                    if isinstance(widget, tk.Text):
                        widget.delete('1.0', 'end')
                        widget.insert('1.0', clean_str)
                    else:
                        widget.delete(0, 'end')
                        widget.insert(0, clean_str)
                    cleaned.append((key, raw, clean_str))
                except Exception:
                    pass
        return cleaned

    def _update_completion_rate(self, event=None):
        """Auto-calculate completion rate from calls completed / total discharges or from missed calls."""
        try:
            total_text = self.get_field_value(self.data.get('total_discharges'))
            calls_text = self.get_field_value(self.data.get('calls_completed'))
            missed_text = self.get_field_value(self.data.get('missed_calls'))

            total = self._parse_number(total_text)
            calls = self._parse_number(calls_text)
            missed = self._parse_number(missed_text)

            # prefer explicit calls_completed if provided; otherwise derive from missed
            if calls is None and missed is not None and total is not None:
                calls = max(0.0, total - missed)

            if total is None or total == 0 or calls is None:
                # clear or leave as-is if not calculable
                return

            pct = (calls / total) * 100.0
            # format with one decimal unless integer
            if abs(pct - int(pct)) < 1e-9:
                pct_str = f"{int(pct)}%"
            else:
                pct_str = f"{pct:.1f}%"

            # set into completion_rate field
            if 'completion_rate' in self.data:
                self.set_field_value(self.data['completion_rate'], pct_str)
                # clear validation marker if present
                if 'completion_rate' in getattr(self, 'validation_markers', {}):
                    self.validation_markers['completion_rate'].config(text='')
        except Exception:
            return

    def show_team_chart(self, info):
        # info is the dict with 'metrics' list containing combo,current,goal widgets
        labels = []
        current_vals = []
        goal_vals = []
        for m in info.get('metrics', []):
            label = m.get('combo').get() if m.get('combo') else ''
            cur = self._parse_number(m.get('current').get()) if m.get('current') else None
            goal = self._parse_number(m.get('goal').get()) if m.get('goal') else None
            if label:
                labels.append(label)
                current_vals.append(cur if cur is not None else 0)
                goal_vals.append(goal if goal is not None else 0)

        if not labels:
            messagebox.showinfo('No Data', 'No metrics available for this team to plot.')
            return

        win = tk.Toplevel(self.root)
        win.title('Team Metrics Chart')
        fig = Figure(figsize=(6,4), dpi=100)
        ax = fig.add_subplot(111)
        x = range(len(labels))
        ax.bar([i-0.2 for i in x], current_vals, width=0.4, label='Current')
        ax.bar([i+0.2 for i in x], goal_vals, width=0.4, label='Goal')
        ax.set_xticks(list(x))
        ax.set_xticklabels(labels, rotation=30, ha='right')
        ax.set_ylabel('Value')
        ax.set_title('Current vs Goal')
        ax.legend()
        fig.tight_layout()

        canvas = FigureCanvasTkAgg(fig, master=win)
        canvas.draw()
        canvas.get_tk_widget().pack(fill='both', expand=True)

    def show_discharge_chart(self):
        # Use completion_rate or calls to build a simple pie chart
        rate_text = self.get_field_value(self.completion_rate)
        rate = self._parse_number(rate_text)
        if rate is None:
            # fallback to compute from calls
            try:
                calls = float(self.get_field_value(self.calls_completed) or 0)
                total = float(self.get_field_value(self.total_discharges) or 0)
                if total > 0:
                    rate = (calls / total) * 100.0
                else:
                    rate = None
            except Exception:
                rate = None

        if rate is None:
            messagebox.showinfo('No Data', 'Completion rate or calls not available to plot.')
            return

        completed = rate
        remaining = max(0, 100.0 - completed)

        win = tk.Toplevel(self.root)
        win.title('48 Hour Discharge Calls Completion')
        fig = Figure(figsize=(5,4), dpi=100)
        ax = fig.add_subplot(111)
        ax.pie([completed, remaining], labels=[f'Completed ({completed:.1f}%)', f'Not Completed ({remaining:.1f}%)'], colors=['#4caf50','#f44336'], autopct='%1.1f%%')
        ax.set_title('48 Hour Call Completion')
        fig.tight_layout()

        canvas = FigureCanvasTkAgg(fig, master=win)
        canvas.draw()
        canvas.get_tk_widget().pack(fill='both', expand=True)
    
    def show_diabetes_chart(self):
        # Plot diabetes-related panel management percentages
        a1c_text = self.get_field_value(getattr(self, 'dm_a1c_updated', None))
        micro_text = self.get_field_value(getattr(self, 'dm_microalbuminuria', None))
        total_text = self.get_field_value(getattr(self, 'dm_total_patients', None))

        a1c = self._parse_number(a1c_text)
        micro = self._parse_number(micro_text)

        labels = []
        vals = []
        if a1c is not None:
            labels.append('A1c up to date')
            vals.append(a1c)
        if micro is not None:
            labels.append('Microalbuminuria up to date')
            vals.append(micro)

        if not labels:
            messagebox.showinfo('No Data', 'No diabetes panel percentages available to plot.')
            return

        win = tk.Toplevel(self.root)
        win.title('Diabetes Panel Management')
        fig = Figure(figsize=(6,4), dpi=100)
        ax = fig.add_subplot(111)
        x = range(len(labels))
        ax.bar(x, vals, color=['#42a5f5', '#66bb6a'])
        ax.set_xticks(list(x))
        ax.set_xticklabels(labels, rotation=20, ha='right')
        ax.set_ylabel('Percent (%)')
        title = 'Diabetes Panel Management'
        if total_text:
            title += f" — Panel size: {self._format_number_for_display(total_text)}"
        ax.set_title(title)
        fig.tight_layout()

        canvas = FigureCanvasTkAgg(fig, master=win)
        canvas.draw()
        canvas.get_tk_widget().pack(fill='both', expand=True)

    def show_red_teams_summary_chart(self):
        # Aggregate red metrics per team (count where current < goal)
        team_names = []
        red_counts = []
        for _, info in getattr(self, 'red_teams', []):
            # team name may be stored as Entry widget in info['name']
            name_widget = info.get('name')
            try:
                team_name = name_widget.get() if name_widget is not None else ''
            except Exception:
                team_name = str(name_widget) or ''

            count = 0
            for m in info.get('metrics', []):
                cur_w = m.get('current')
                goal_w = m.get('goal')
                cur = self._parse_number(cur_w.get()) if cur_w else None
                goal = self._parse_number(goal_w.get()) if goal_w else None
                if cur is not None and goal is not None:
                    try:
                        if cur < goal:
                            count += 1
                    except Exception:
                        pass
            if team_name.strip() == '':
                team_name = f'Team {len(team_names)+1}'
            team_names.append(team_name)
            red_counts.append(count)

        if not team_names:
            messagebox.showinfo('No Teams', 'No red teams have been added yet.')
            return

        # if all counts are zero, inform user
        if all(c == 0 for c in red_counts):
            messagebox.showinfo('No Red Metrics', 'No red metrics found across teams to summarize.')
            return

        win = tk.Toplevel(self.root)
        win.title('Red Teams Summary')
        fig = Figure(figsize=(7,4), dpi=100)
        ax = fig.add_subplot(111)
        x = range(len(team_names))
        ax.bar(x, red_counts, color='#ef5350')
        ax.set_xticks(list(x))
        ax.set_xticklabels(team_names, rotation=30, ha='right')
        ax.set_ylabel('Count of Red Metrics')
        ax.set_title('Red Metrics by Team')
        fig.tight_layout()

        canvas = FigureCanvasTkAgg(fig, master=win)
        canvas.draw()
        canvas.get_tk_widget().pack(fill='both', expand=True)
    
    def create_monthly_tab(self, notebook):
        frame = ttk.Frame(notebook)
        notebook.add(frame, text="Monthly")
        
        canvas = tk.Canvas(frame)
        scrollbar = tk.Scrollbar(frame, orient="vertical", command=canvas.yview, width=30, bg='#4a4a4a', troughcolor='#ffffff', activebackground='#2f2f2f')
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Panel Management - Diabetes
        diabetes_frame = self.create_form_section(scrollable_frame, "Monthly Panel Management - Diabetes")
        self.dm_total_patients = self.add_text_field(diabetes_frame, "Total DM Patients on Panel (last month):", "dm_total_patients")
        self.dm_a1c_updated = self.add_text_field(diabetes_frame, "% with A1c up to date (last month):", "dm_a1c_updated")
        self.dm_microalbuminuria = self.add_text_field(diabetes_frame, "% with Microalbuminuria/Urine Albumin Creatinine Ratio up to date:", "dm_microalbuminuria")
        ttk.Button(diabetes_frame, text="Show Diabetes Chart", command=self.show_diabetes_chart).grid(row=len(diabetes_frame.grid_slaves()), column=1, sticky='e', padx=8, pady=6)
        
        # Panel Management - Data in the Red
        red_frame = self.create_form_section(scrollable_frame, "Monthly Panel Management - Data in the Red")
        self.total_pact_teams = self.add_text_field(red_frame, "Total PACT Teams:", "total_pact_teams")
        self.red_teams_count = self.add_text_field(red_frame, "Number of PACT Teams with Any Red Metric:", "red_teams_count")
        ttk.Button(red_frame, text="Show Red Teams Summary", command=self.show_red_teams_summary_chart).grid(row=len(red_frame.grid_slaves()), column=1, sticky='e', padx=8, pady=6)
        
        # Top 3 Red Teams
        # Red Teams (dynamic list) moved here from Support tab
        red_section = self.create_form_section(scrollable_frame, "Monthly Panel Management - Teams with Red Data Metrics")
        red_btn_frame = ttk.Frame(red_section)
        red_btn_frame.pack(fill="x")
        ttk.Button(red_btn_frame, text="Add Team with Red Metrics", command=self.add_red_team).pack(side="right")
        self.red_container = ttk.Frame(red_section)
        self.red_container.pack(fill="x", pady=5)
        
        # Chart Audits
        audits_frame = self.create_form_section(scrollable_frame, "Monthly Nursing Chart Audits & Outliers")
        self.charts_audited = self.add_text_field(audits_frame, "Number of Charts Audited This Month:", "charts_audited")
        self.outliers_count = self.add_text_field(audits_frame, "Number of Outliers:", "outliers_count")
        
        # Outlier reasons
        ttk.Label(audits_frame, text="Outlier Reasons (check all that apply):").grid(sticky="nw", pady=(10,5))
        self.outlier_incomplete = self.add_checkbox(audits_frame, "Incomplete assessment", "outlier_incomplete")
        self.outlier_missing_education = self.add_checkbox(audits_frame, "Missing patient education", "outlier_missing_education")
        self.outlier_incorrect_template = self.add_checkbox(audits_frame, "Incorrect/old template", "outlier_incorrect_template")
        self.outlier_delayed_doc = self.add_checkbox(audits_frame, "Delayed documentation (>24 hrs)", "outlier_delayed_doc")
        self.outlier_protocol = self.add_checkbox(audits_frame, "Protocol not followed", "outlier_protocol")
        self.outlier_other = self.add_text_field(audits_frame, "Other:", "outlier_other")
        
        # Response to outliers
        response_frame = self.create_form_section(scrollable_frame, "Response to Outliers")
        self.response_feedback = self.add_checkbox(response_frame, "1:1 feedback", "response_feedback")
        self.response_group_education = self.add_checkbox(response_frame, "Group education planned/updated", "response_group_education")
        self.response_workflow = self.add_checkbox(response_frame, "Workflow clarified/updated", "response_workflow")
        self.response_policy = self.add_checkbox(response_frame, "Policy/procedure review requested", "response_policy")
        self.response_pdsa = self.add_checkbox(response_frame, "Added to PDSA focus", "response_pdsa")
        self.response_other = self.add_text_field(response_frame, "Other:", "response_other")

        # Repeated outliers
        repeated_frame = self.create_form_section(scrollable_frame, "Repeated Outliers")
        self.repeated_staff = self.add_text_field(repeated_frame, "Staff Initials with Repeated Outliers:", "repeated_staff")
        self.repeated_plan = self.add_text_field(repeated_frame, "Plan:", "repeated_plan", height=2)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        return frame
    
    def create_pdsa_tab(self, notebook):
        frame = ttk.Frame(notebook)
        notebook.add(frame, text="PDSA & Education/Coaching")
        
        canvas = tk.Canvas(frame)
        scrollbar = tk.Scrollbar(frame, orient="vertical", command=canvas.yview, width=30, bg='#4a4a4a', troughcolor='#ffffff', activebackground='#2f2f2f')
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # Noncompliance & Discipline (dynamic list)
        discipline_section = self.create_form_section(scrollable_frame, "Improvement Opportunity & Training")
        btn_frame = ttk.Frame(discipline_section)
        btn_frame.pack(fill="x")
        ttk.Button(btn_frame, text="Add Issue", command=self.add_discipline_issue).pack(side="right")
        self.discipline_container = ttk.Frame(discipline_section)
        self.discipline_container.pack(fill="x", pady=5)

        # PDSA Cycles & Projects (dynamic list)
        pdsa_section = self.create_form_section(scrollable_frame, "PDSA Cycles & Projects")
        pdsa_btn_frame = ttk.Frame(pdsa_section)
        pdsa_btn_frame.pack(fill="x")
        ttk.Button(pdsa_btn_frame, text="Add Project", command=self.add_pdsa_project).pack(side="right")
        self.pdsa_container = ttk.Frame(pdsa_section)
        self.pdsa_container.pack(fill="x", pady=5)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        return frame
    
    def create_support_tab(self, notebook):
        frame = ttk.Frame(notebook)
        notebook.add(frame, text="Support & Barriers")
        
        canvas = tk.Canvas(frame)
        scrollbar = tk.Scrollbar(frame, orient="vertical", command=canvas.yview, width=30, bg='#4a4a4a', troughcolor='#ffffff', activebackground='#2f2f2f')
        scrollable_frame = ttk.Frame(canvas)
        
        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )
        
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)
        
        # What do you need from me?
        needs_frame = self.create_form_section(scrollable_frame, "Actions and Support Needed")
        self.barriers_1 = self.add_text_field(needs_frame, "Top Barrier #1:", "barriers_1")
        self.barriers_2 = self.add_text_field(needs_frame, "Top Barrier #2:", "barriers_2")
        self.barriers_3 = self.add_text_field(needs_frame, "Top Barrier #3:", "barriers_3")
        
        # Support requested
        support_frame = self.create_form_section(scrollable_frame, "Support Requested (check all that apply)")
        self.support_data_access = self.add_checkbox(support_frame, "Data/report access or clarification", "support_data_access")
        self.support_it = self.add_checkbox(support_frame, "IT build/template change", "support_it")
        self.support_hr = self.add_checkbox(support_frame, "HR/Employee and labor relations support", "support_hr")
        self.support_provider = self.add_checkbox(support_frame, "Provider leadership engagement", "support_provider")
        self.support_education = self.add_checkbox(support_frame, "Education/resources for staff", "support_education")
        self.support_staffing = self.add_checkbox(support_frame, "Additional staffing/float support", "support_staffing")
        self.support_other = self.add_text_field(support_frame, "Other:", "support_other")
        
        # Risk assessment
        risk_frame = self.create_form_section(scrollable_frame, "Risk Assessment")
        ttk.Label(risk_frame, text="Potential Risks").grid(row=0, column=0, columnspan=2, sticky="w", pady=5)
        self.risk_no = tk.BooleanVar()
        self.risk_yes = tk.BooleanVar()
        ttk.Radiobutton(risk_frame, text="No", variable=self.risk_no, command=lambda: self.risk_yes.set(False)).grid(row=1, column=0, sticky="w", padx=10)
        ttk.Radiobutton(risk_frame, text="Yes", variable=self.risk_yes, command=lambda: self.risk_no.set(False)).grid(row=1, column=1, sticky="w", padx=10)
        self.risk_describe = self.add_text_field(risk_frame, "Describe Risk:", "risk_describe", height=3)
        
        # Team metrics selection for red areas
        # (metrics dropdown moved to Monthly tab per-team)
        
        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        
        return frame
    
    def create_form_section(self, parent, title):
        """Create a labeled section for form fields"""
        section = ttk.LabelFrame(parent, text=title, padding=12)
        section.pack(fill="x", padx=8, pady=10)
        try:
            section.columnconfigure(1, weight=1)
        except Exception:
            pass
        return section
    
    def add_text_field(self, parent, label, key, height=1, width=50):
        """Add a text field to a section"""
        # Get the next row number
        row = len([w for w in parent.winfo_children()])
        
        if label:
            ttk.Label(parent, text=label).grid(row=row, column=0, sticky="ne", padx=8, pady=6)
        
        if height == 1:
            if self.default_font:
                entry = ttk.Entry(parent, width=width)
                entry.configure(font=(self.font_name, self.font_size))
            else:
                entry = ttk.Entry(parent, width=width)
        else:
            if self.default_font:
                entry = tk.Text(parent, height=height, width=width, font=(self.font_name, self.font_size))
            else:
                entry = tk.Text(parent, height=height, width=width)
        
        entry.grid(row=row, column=1, sticky="ew", padx=8, pady=6)
        entry.key = key
        self.data[key] = entry
        # If this field should be numeric, add a validation marker and bind focus-out
        if key in getattr(self, 'numeric_keys', []):
            marker = tk.Label(parent, text='', fg='red')
            try:
                marker.configure(font=(self.font_name, max(self.font_size-2, 10)))
            except Exception:
                pass
            marker.grid(row=row, column=2, sticky='w', padx=4)
            self.validation_markers[key] = marker

            def on_focus_out(event, k=key, w=entry, m=marker):
                val = self.get_field_value(w)
                if val and self._parse_number(val) is None:
                    m.config(text='✱ invalid number')
                else:
                    m.config(text='')

            # bind focus-out for both Entry and Text
            try:
                entry.bind('<FocusOut>', on_focus_out)
            except Exception:
                pass
            # update completion rate when discharge/calls/missed change
            if key in ('total_discharges', 'calls_completed', 'missed_calls'):
                try:
                    entry.bind('<FocusOut>', lambda e: self._update_completion_rate())
                except Exception:
                    pass
        return entry
    
    def add_checkbox(self, parent, label, key):
        """Add a checkbox to a section"""
        # Get the next row number
        row = len([w for w in parent.winfo_children()])
        
        var = tk.BooleanVar()
        check = ttk.Checkbutton(parent, text=label, variable=var)
        if self.default_font:
            try:
                check.configure(font=(self.font_name, self.font_size))
            except Exception:
                pass
        check.grid(row=row, column=0, columnspan=2, sticky="w", padx=8, pady=6)
        check.key = key
        self.data[key] = var
        return var
    
    def get_field_value(self, field):
        """Get value from any field type"""
        if isinstance(field, tk.BooleanVar) or isinstance(field, tk.StringVar):
            return field.get()
        elif isinstance(field, tk.Text):
            return field.get("1.0", "end-1c")
        else:  # ttk.Entry
            return field.get()
    
    def set_field_value(self, field, value):
        """Set value for any field type"""
        if isinstance(field, tk.BooleanVar) or isinstance(field, tk.StringVar):
            field.set(value)
        elif isinstance(field, tk.Text):
            field.delete("1.0", "end")
            field.insert("1.0", value)
        else:  # ttk.Entry
            field.delete(0, "end")
            field.insert(0, value)
    
    def save_form(self):
        """Save form data"""
        if not self.current_file:
            self.save_form_as()
            return
        
        self.save_data_to_file(self.current_file)
        messagebox.showinfo("Success", "Form saved successfully!")
    
    def save_form_as(self):
        """Save form as new file"""
        file_path = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")],
            initialdir="c:/Users/FreeP/Desktop/manager folder"
        )
        
        if file_path:
            self.current_file = file_path
            self.save_form()
    
    def save_data_to_file(self, file_path):
        """Save all form data to JSON file"""
        # validate numeric inputs before saving
        issues = self._collect_numeric_issues()
        if issues:
            msg = "Some fields expected numeric values appear invalid.\n\n"
            msg += "Examples:\n"
            for k, w, v in issues[:6]:
                msg += f" - {k}: {v}\n"
            msg += "\nWould you like the app to try to auto-clean these fields (strip % and commas)?"
            if messagebox.askyesno("Numeric Validation", msg):
                cleaned = self._attempt_autoclean_numeric_issues(issues)
                if cleaned:
                    messagebox.showinfo("Auto-clean", f"Cleaned {len(cleaned)} values.")
                else:
                    messagebox.showwarning("Auto-clean", "Could not clean the detected fields. Please correct them manually.")
            else:
                # user chose not to autoclean; abort save so they can fix
                if not messagebox.askyesno("Abort Save?", "Do you want to cancel save and correct values? (No will continue saving raw values)"):
                    pass

        data_dict = {}
        for key, field in self.data.items():
            data_dict[key] = self.get_field_value(field)

        # Discipline issues
        data_dict['discipline_issues'] = []
        for frame, fields in self.discipline_issues:
            issue = {}
            for k, w in fields.items():
                if isinstance(w, tk.Text):
                    issue[k] = w.get("1.0", "end-1c")
                else:
                    try:
                        issue[k] = w.get()
                    except Exception:
                        issue[k] = ""
            data_dict['discipline_issues'].append(issue)

        # PDSA projects
        data_dict['pdsa_projects'] = []
        for frame, fields in self.pdsa_projects:
            proj = {}
            for k, w in fields.items():
                if isinstance(w, tk.Text):
                    proj[k] = w.get("1.0", "end-1c")
                else:
                    try:
                        proj[k] = w.get()
                    except Exception:
                        proj[k] = ""
            data_dict['pdsa_projects'].append(proj)

        # Red teams
        data_dict['red_teams'] = []
        for frame, info in self.red_teams:
            team = {}
            team['name'] = info.get('name').get() if info.get('name') else ''
            # metrics
            team['metrics'] = []
            for m in info.get('metrics', []):
                metric = {}
                metric['metric'] = m['combo'].get() if m.get('combo') else ''
                metric['current'] = m['current'].get() if m.get('current') else ''
                metric['goal'] = m['goal'].get() if m.get('goal') else ''
                team['metrics'].append(metric)
            # barriers
            team['barriers'] = []
            for b in info.get('barriers', []):
                team['barriers'].append({
                    'name': b['label'],
                    'checked': bool(b['var'].get()),
                    'support': b['support'].get() if b.get('support') else ''
                })
            data_dict['red_teams'].append(team)

        with open(file_path, 'w') as f:
            json.dump(data_dict, f, indent=4)
    
    def open_form(self):
        """Open saved form data"""
        file_path = filedialog.askopenfilename(
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")],
            initialdir="c:/Users/FreeP/Desktop/manager folder"
        )
        
        if file_path:
            self.current_file = file_path
            with open(file_path, 'r') as f:
                data_dict = json.load(f)
            
            # Load static fields
            for key, value in data_dict.items():
                if key in self.data and key not in ('discipline_issues', 'pdsa_projects'):
                    self.set_field_value(self.data[key], value)

            # Clear existing dynamic entries
            for frm, _ in list(self.discipline_issues):
                try:
                    frm.destroy()
                except Exception:
                    pass
            self.discipline_issues = []

            for frm, _ in list(self.pdsa_projects):
                try:
                    frm.destroy()
                except Exception:
                    pass
            self.pdsa_projects = []

            # clear red teams
            for frm, _ in list(self.red_teams):
                try:
                    frm.destroy()
                except Exception:
                    pass
            self.red_teams = []

            # Load discipline issues
            for issue in data_dict.get('discipline_issues', []):
                self.add_discipline_issue(prefill=issue)

            # Load pdsa projects
            for proj in data_dict.get('pdsa_projects', []):
                self.add_pdsa_project(prefill=proj)

            # Load red teams
            for frm, _ in list(self.red_teams):
                try:
                    frm.destroy()
                except Exception:
                    pass
            self.red_teams = []
            for team in data_dict.get('red_teams', []):
                self.add_red_team(prefill=team)

            messagebox.showinfo("Success", "Form loaded successfully!")
    
    def new_form(self):
        """Clear all form fields"""
        if messagebox.askyesno("New Form", "Clear all fields and start a new form?"):
            for field in self.data.values():
                if isinstance(field, tk.BooleanVar) or isinstance(field, tk.StringVar):
                    field.set("")
                elif isinstance(field, tk.Text):
                    field.delete("1.0", "end")
                else:  # ttk.Entry
                    field.delete(0, "end")
            # clear dynamic lists
            for frm, _ in list(self.discipline_issues):
                try:
                    frm.destroy()
                except Exception:
                    pass
            self.discipline_issues = []

            for frm, _ in list(self.pdsa_projects):
                try:
                    frm.destroy()
                except Exception:
                    pass
            self.pdsa_projects = []

            for frm, _ in list(self.red_teams):
                try:
                    frm.destroy()
                except Exception:
                    pass
            self.red_teams = []

            self.current_file = None
            messagebox.showinfo("Success", "New form started!")
    
    def export_to_word(self):
        """Export form data to Word document"""
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
            initialdir="c:/Users/FreeP/Desktop/manager folder"
        )
        
        if not file_path:
            return
        doc = Document()
        doc.styles['Normal'].font.name = self.font_name
        doc.styles['Normal'].font.size = Pt(self.font_size)
        doc.add_heading('Nurse Manager Huddle Report', 0)
        doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph()

        # Summary / Huddle info
        doc.add_heading('Huddle Information', level=1)
        for key, label in [
            ("service_unit", "Service/Unit"),
            ("nurse_manager", "Nurse Manager"),
            ("week_of", "Week of"),
            ("huddle_date", "Huddle Date"),
            ("participants", "Participants")
        ]:
            if key in self.data:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(str(self.get_field_value(self.data[key])))

        # Discharge summary
        doc.add_heading('48 Hour Discharge Calls', level=1)
        for key, label in [
            ("total_discharges", "Total Discharges"),
            ("calls_completed", "Calls Completed"),
            ("completion_rate", "Completion Rate"),
            ("missed_calls", "Missed Calls")
        ]:
            if key in self.data:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(str(self.get_field_value(self.data[key])))

        # Discipline issues
        doc.add_heading('Noncompliance & Discipline', level=1)
        if not self.discipline_issues:
            doc.add_paragraph('No discipline issues recorded.')
        else:
            for i, (frm, fields) in enumerate(self.discipline_issues, start=1):
                doc.add_heading(f'Issue #{i}', level=2)
                for k, w in fields.items():
                    val = w.get('1.0', 'end-1c') if isinstance(w, tk.Text) else w.get()
                    p = doc.add_paragraph()
                    p.add_run(f"{k.replace('_',' ').title()}: ").bold = True
                    p.add_run(val)

        # PDSA projects
        doc.add_heading('PDSA Projects & Tests', level=1)
        if not self.pdsa_projects:
            doc.add_paragraph('No PDSA projects recorded.')
        else:
            for i, (frm, fields) in enumerate(self.pdsa_projects, start=1):
                name = fields.get('name').get() if fields.get('name') else ''
                doc.add_heading(f'Project #{i}: {name}', level=2)
                for k in ('focus','phase'):
                    if k in fields:
                        val = fields[k].get() if not isinstance(fields[k], tk.Text) else fields[k].get('1.0','end-1c')
                        p = doc.add_paragraph()
                        p.add_run(f"{k.replace('_',' ').title()}: ").bold = True
                        p.add_run(val)
                # metrics table
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Light List Accent 1'
                hdr = table.rows[0].cells
                hdr[0].text = 'Metric'
                hdr[1].text = 'Value'
                hdr[2].text = 'Goal'
                for j in range(1,4):
                    mk = fields.get(f'metric{j}')
                    mv = fields.get(f'metric{j}_value')
                    mg = fields.get(f'metric{j}_goal')
                    if mk:
                        metric = mk.get() if not isinstance(mk, tk.Text) else mk.get('1.0','end-1c')
                        value = mv.get() if mv and not isinstance(mv, tk.Text) else (mv.get('1.0','end-1c') if mv else '')
                        goal = mg.get() if mg and not isinstance(mg, tk.Text) else (mg.get('1.0','end-1c') if mg else '')
                        if metric.strip() or value.strip() or goal.strip():
                            row = table.add_row().cells
                            row[0].text = metric
                            row[1].text = value
                            row[2].text = goal
                # learning
                if 'learning' in fields:
                    p = doc.add_paragraph()
                    p.add_run('Key Learning: ').bold = True
                    p.add_run(fields['learning'].get('1.0','end-1c') if isinstance(fields['learning'], tk.Text) else fields['learning'].get())

        # Red teams
        doc.add_heading('Red Teams (Top Areas in the Red)', level=1)
        if not self.red_teams:
            doc.add_paragraph('No red teams recorded.')
        else:
            for i, (frm, info) in enumerate(self.red_teams, start=1):
                team_name = info.get('name').get() if info.get('name') else f'Red Team {i}'
                doc.add_heading(f'{team_name}', level=2)
                # metrics table
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Light List Accent 2'
                hdr = table.rows[0].cells
                hdr[0].text = 'Metric'
                hdr[1].text = 'Current'
                hdr[2].text = 'Goal'
                for m in info.get('metrics', []):
                    metric = m.get('combo').get() if m.get('combo') else ''
                    cur = m.get('current').get() if m.get('current') else ''
                    goal = m.get('goal').get() if m.get('goal') else ''
                    if metric or cur or goal:
                        row = table.add_row().cells
                        row[0].text = metric
                        row[1].text = cur
                        row[2].text = goal
                # barriers
                doc.add_paragraph('Barriers and Support:').bold = True
                for b in info.get('barriers', []):
                    checked = b['var'].get()
                    support = b['support'].get() if b.get('support') else ''
                    p = doc.add_paragraph()
                    p.add_run(f"- {b['label']}: ").bold = True
                    p.add_run(('Yes' if checked else 'No') + (f"; Support: {support}" if support else ''))

        # final save
        doc.save(file_path)
        messagebox.showinfo("Success", f"Form exported to Word successfully!\n{file_path}")
    
    def export_to_excel(self):
        """Export form data to Excel spreadsheet"""
        file_path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")],
            initialdir="c:/Users/FreeP/Desktop/manager folder"
        )
        
        if not file_path:
            return
        
        wb = openpyxl.Workbook()
        # Summary sheet
        ws = wb.active
        ws.title = "Summary"
        ws.append(['Field', 'Value'])
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        for cell in ('A1', 'B1'):
            ws[cell].font = header_font
            ws[cell].fill = header_fill

        for key, field in self.data.items():
            ws.append([key.replace('_',' ').title(), str(self.get_field_value(field))])
        ws.column_dimensions['A'].width = 40
        ws.column_dimensions['B'].width = 60

        # Discipline issues sheet
        ds = wb.create_sheet('Discipline')
        ds.append(['Issue #', 'Field', 'Value'])
        ds['A1'].font = header_font
        ds['A1'].fill = header_fill
        ds['B1'].font = header_font
        ds['B1'].fill = header_fill
        ds['C1'].font = header_font
        ds['C1'].fill = header_fill
        for i, (frm, fields) in enumerate(self.discipline_issues, start=1):
            for k, w in fields.items():
                val = w.get('1.0', 'end-1c') if isinstance(w, tk.Text) else w.get()
                ds.append([i, k.replace('_',' ').title(), val])

        # PDSA projects
        ps = wb.create_sheet('PDSA')
        ps.append(['Project #', 'Name', 'Focus', 'Phase', 'Metric', 'Value', 'Goal'])
        for i, (frm, fields) in enumerate(self.pdsa_projects, start=1):
            name = fields.get('name').get() if fields.get('name') else ''
            focus = fields.get('focus').get() if fields.get('focus') else ''
            phase = fields.get('phase').get() if fields.get('phase') else ''
            for j in range(1,4):
                mk = fields.get(f'metric{j}')
                mv = fields.get(f'metric{j}_value')
                mg = fields.get(f'metric{j}_goal')
                metric = mk.get() if mk and not isinstance(mk, tk.Text) else (mk.get('1.0','end-1c') if mk else '')
                val = mv.get() if mv and not isinstance(mv, tk.Text) else (mv.get('1.0','end-1c') if mv else '')
                goal = mg.get() if mg and not isinstance(mg, tk.Text) else (mg.get('1.0','end-1c') if mg else '')
                if metric.strip() or val.strip() or goal.strip():
                    ps.append([i, name, focus, phase, metric, val, goal])

        # Red teams
        rs = wb.create_sheet('Red Teams')
        rs.append(['Team', 'Metric', 'Current', 'Goal', 'Barrier', 'Barrier Support'])
        for i, (frm, info) in enumerate(self.red_teams, start=1):
            team = info.get('name').get() if info.get('name') else f'Team {i}'
            # metrics rows
            for m in info.get('metrics', []):
                metric = m.get('combo').get() if m.get('combo') else ''
                cur = m.get('current').get() if m.get('current') else ''
                goal = m.get('goal').get() if m.get('goal') else ''
                # empty barrier placeholders for metric rows
                rs.append([team, metric, cur, goal, '', ''])
            # barriers rows
            for b in info.get('barriers', []):
                rs.append([team, '', '', '', b['label'], b.get('support').get() if b.get('support') else ''])

        # style column widths
        for sheet in (ds, ps, rs):
            sheet.column_dimensions['A'].width = 20
            sheet.column_dimensions['B'].width = 30
            sheet.column_dimensions['C'].width = 20
        wb.save(file_path)
        messagebox.showinfo("Success", f"Form exported to Excel successfully!\n{file_path}")
    
    def show_about(self):
        messagebox.showinfo("About", " Meeting and Huddle Tool\n-----------------------------\nA comprehensive GUI for \nmanaging weekly and monthly huddle reports.\nSuggestions or changes are welcome!\nTMT_labs Version 1.0")

if __name__ == "__main__":
    root = tk.Tk()
    app = NurseManagerApp(root)
    root.mainloop()
